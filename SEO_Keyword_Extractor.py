import customtkinter as ctk
from tkinter import messagebox
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import urllib.parse
import os
import json
from datetime import datetime
import threading
import re
from pathlib import Path
import pandas as pd
from docx import Document

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


# ================= BACKEND =================
class SEODataExtractor:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0"
        })

    def generate_urls_for_keyword(self, keyword):
        q = urllib.parse.quote(keyword)
        urls = [
            f"https://en.wikipedia.org/wiki/{keyword.replace(' ','_')}",
            f"https://medium.com/search?q={q}",
            f"https://quotes.toscrape.com/",
            f"https://example.com",
        ]
        return [{'title': f'{keyword} - {urlparse(u).netloc}', 'url': u} for u in urls]

    def extract_page_data(self, url):
        try:
            r = self.session.get(url, timeout=10)
            soup = BeautifulSoup(r.text, "html.parser")

            images = self.extract_images(soup, url)

            return {
                "url": url,
                "title": soup.title.string if soup.title else "No title",
                "description": "",
                "keywords": "",
                "images": images,
                "images_count": len(images),
                "status": "success",
                "scraped_at": datetime.now().isoformat()
            }
        except:
            return {"url": url, "status": "error"}

    def extract_images(self, soup, base):
        imgs = []
        for img in soup.find_all("img"):
            src = img.get("src")
            if src:
                full = urljoin(base, src)
                if any(full.endswith(x) for x in ['.jpg','.png','.jpeg','.webp']):
                    imgs.append(full)
        return imgs

    def download_images(self, imgs, folder):
        os.makedirs(folder, exist_ok=True)
        for i, url in enumerate(imgs[:30]):
            try:
                r = self.session.get(url, stream=True)
                with open(os.path.join(folder, f"img_{i}.jpg"), "wb") as f:
                    for chunk in r.iter_content(1024):
                        f.write(chunk)
            except:
                pass

    def save(self, keyword, data, fmt, download):
        folder = Path(f"SEO_{keyword}_{datetime.now().strftime('%H%M%S')}")
        folder.mkdir(exist_ok=True)

        all_imgs = []
        for p in data["pages"]:
            all_imgs += p.get("images", [])

        if download:
            self.download_images(all_imgs, folder/"images")

        df = pd.DataFrame(data["pages"])

        if fmt == "json":
            with open(folder/"data.json","w") as f:
                json.dump(data,f,indent=2)

        elif fmt == "csv":
            df.to_csv(folder/"data.csv", index=False)

        elif fmt == "excel":
            df.to_excel(folder/"data.xlsx", index=False)

        elif fmt == "docx":
            doc = Document()
            doc.add_heading(keyword,0)
            for _,row in df.iterrows():
                doc.add_paragraph(str(row))
            doc.save(folder/"data.docx")

        return str(folder)


# ================= UI =================
class App:
    def __init__(self):
        self.root = ctk.CTk()
        self.root.geometry("1100x800")
        self.root.title("SEO Data Extractor Pro v3.0")

        self.extractor = SEODataExtractor()

        self.build_ui()

    def build_ui(self):
        # Title
        ctk.CTkLabel(self.root,
                     text="🔍 SEO Data Extractor Pro",
                     font=ctk.CTkFont(size=32, weight="bold")
                     ).pack(pady=20)

        main = ctk.CTkFrame(self.root)
        main.pack(fill="both", expand=True, padx=40, pady=20)

        # Keyword
        ctk.CTkLabel(main, text="Keyword:", font=("Arial",18)).pack(anchor="w", padx=20, pady=(20,5))
        self.entry = ctk.CTkEntry(main, height=45, placeholder_text="Enter SEO keyword...")
        self.entry.pack(fill="x", padx=20)

        # Format
        ctk.CTkLabel(main, text="📄 Export Format:", font=("Arial",18,"bold")).pack(anchor="w", padx=20, pady=15)

        self.format = ctk.StringVar(value="json")

        row = ctk.CTkFrame(main)
        row.pack(fill="x", padx=20)

        for t,v in [("JSON","json"),("CSV","csv"),("Excel","excel"),("Word DOCX","docx")]:
            ctk.CTkRadioButton(row, text=t, variable=self.format, value=v).pack(side="left", padx=10)

        # Checkbox
        self.img_check = ctk.CTkCheckBox(main, text="💾 Download Images")
        self.img_check.pack(anchor="w", padx=20, pady=10)

        # Buttons
        btn_frame = ctk.CTkFrame(main)
        btn_frame.pack(pady=15)

        self.start_btn = ctk.CTkButton(btn_frame, text="🚀 START EXTRACTION",
                                      fg_color="#10b981",
                                      height=50,
                                      command=self.start)
        self.start_btn.pack(side="left", padx=10)

        self.open_btn = ctk.CTkButton(btn_frame, text="📁 OPEN FOLDER",
                                     state="disabled",
                                     command=self.open_folder)
        self.open_btn.pack(side="left", padx=10)

        # Progress
        self.progress = ctk.CTkProgressBar(main, height=25)
        self.progress.pack(fill="x", padx=20, pady=10)

        self.status = ctk.CTkLabel(main, text="✅ Ready to extract!")
        self.status.pack()

        # Log
        self.log = ctk.CTkTextbox(main, height=200)
        self.log.pack(fill="both", expand=True, padx=20, pady=20)

    def log_msg(self, msg):
        self.log.insert("end", msg+"\n")
        self.log.see("end")

    def start(self):
        threading.Thread(target=self.run).start()

    def run(self):
        kw = self.entry.get()
        if not kw:
            messagebox.showerror("Error","Enter keyword")
            return

        self.log.delete("1.0","end")

        urls = self.extractor.generate_urls_for_keyword(kw)
        data = {"keyword":kw,"pages":[]}

        for i,u in enumerate(urls):
            self.log_msg(f"[{i+1}] {u['url']}")
            page = self.extractor.extract_page_data(u["url"])
            data["pages"].append(page)
            self.progress.set((i+1)/len(urls))

        folder = self.extractor.save(
            kw, data, self.format.get(), self.img_check.get()
        )

        self.status.configure(text="✅ Done!")
        self.log_msg(f"Saved: {folder}")
        self.open_btn.configure(state="normal")
        self.folder = folder

    def open_folder(self):
        os.startfile(self.folder)

    def run_app(self):
        self.root.mainloop()


if __name__ == "__main__":
    App().run_app()